"""
文档加载器模块

支持多种文档格式的加载：
- .txt / .md: 纯文本文件，直接读取
- .pdf: PDF 文件
- .docx: Word 文档（需要 python-docx）
- URL: 网页内容（httpx + BeautifulSoup）

所有加载器都返回 LangChain 的 Document 对象列表。
"""
from pathlib import Path
from langchain_core.documents import Document
from bs4 import BeautifulSoup
import httpx
from agent.tools.url_security import validate_public_http_url


def _normalize_text(text: str) -> str:
    """清理网页提取文本，去掉连续空行和完全重复的相邻行。"""
    lines = []
    last = None
    for raw in text.splitlines():
        line = raw.strip()
        if not line or line == last:
            continue
        lines.append(line)
        last = line
    return "\n".join(lines)


def load_file(file_path: str) -> list[Document]:
    """
    加载本地文件

    根据文件扩展名选择不同的解析方式。

    Args:
        file_path: 文件路径

    Returns:
        Document 列表（通常只有一个元素）

    Raises:
        FileNotFoundError: 文件不存在时抛出
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    content = path.read_text(encoding="utf-8")

    # 元数据：记录文件来源，用于后续展示和管理
    metadata = {"source": str(path), "filename": path.name, "type": suffix}

    if suffix in (".txt", ".md"):
        # 纯文本和 Markdown 直接读取
        return [Document(page_content=content, metadata=metadata)]
    elif suffix == ".pdf":
        return _load_pdf(content, metadata)
    elif suffix == ".docx":
        return _load_docx(str(path), metadata)
    else:
        # 其他格式尝试当纯文本读取
        return [Document(page_content=content, metadata=metadata)]


def load_url(url: str) -> list[Document]:
    """
    加载网页内容

    使用 httpx 抓取网页，BeautifulSoup 提取正文。

    Args:
        url: 网页 URL

    Returns:
        Document 列表
    """
    safe_url = validate_public_http_url(url)
    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                      "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
    }
    resp = httpx.get(safe_url, headers=headers, follow_redirects=True, timeout=15)
    validate_public_http_url(str(resp.url))
    resp.raise_for_status()

    # 解析 HTML，去除非内容标签。这里抓取的是服务端返回的 HTML；
    # 对纯前端渲染站点，静态 HTML 里可能没有完整业务内容。
    soup = BeautifulSoup(resp.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()
    main_node = soup.find("main") or soup.find("article") or soup.body or soup
    text = _normalize_text(main_node.get_text(separator="\n", strip=True))
    title = soup.title.string.strip() if soup.title and soup.title.string else url
    is_sparse = len(text) < 800
    extraction_note = ""
    if is_sparse:
        extraction_note = "静态 HTML 内容较少，该页面可能依赖前端渲染，当前未执行浏览器渲染抓取。"

    return [
        Document(
            page_content=text,
            metadata={
                "source": url,
                "type": "url",
                "title": title[:200],
                "content_length": len(text),
                "extraction_quality": "sparse" if is_sparse else "static_html",
                "extraction_note": extraction_note,
            },
        )
    ]


def _load_pdf(content: str, metadata: dict) -> list[Document]:
    """
    加载 PDF 文件（简化实现）

    注意：这里简化为直接读取文本内容。
    完整实现应使用 PyPDF2 或 pdfplumber 提取文本。
    """
    return [Document(page_content=content, metadata=metadata)]


def _load_docx(file_path: str, metadata: dict) -> list[Document]:
    """
    加载 Word 文档

    使用 python-docx 库解析 .docx 文件。
    如果 python-docx 未安装，返回错误提示。

    Args:
        file_path: .docx 文件路径
        metadata: 元数据字典

    Returns:
        Document 列表
    """
    try:
        import docx
        doc = docx.Document(file_path)
        # 提取所有非空段落的文本
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata=metadata)]
    except ImportError:
        return [Document(page_content="python-docx 未安装，无法解析 .docx 文件", metadata=metadata)]
